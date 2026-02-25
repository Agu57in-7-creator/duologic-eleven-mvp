import streamlit as st
from google import genai
from datetime import date, datetime
import json
import os
import pandas as pd
import io
from docx import Document

# --- RUTAS DE LA BASE DE DATOS LOCAL ---
ARCHIVO_PEDIDOS = "pedidos.json"
ARCHIVO_PROVEEDORES = "proveedores.json"
ARCHIVO_ENTREGADOS = "entregados.json"
CARPETA_PROVEEDORES = "archivos_proveedores"
CARPETA_PEDIDOS = "archivos_pedidos"

if not os.path.exists(CARPETA_PROVEEDORES): os.makedirs(CARPETA_PROVEEDORES)
if not os.path.exists(CARPETA_PEDIDOS): os.makedirs(CARPETA_PEDIDOS)

# --- FUNCIONES DE BASE DE DATOS ---
def cargar_datos():
    if os.path.exists(ARCHIVO_PEDIDOS):
        with open(ARCHIVO_PEDIDOS, "r") as f:
            try:
                pedidos = json.load(f)
                for p in pedidos:
                    p["fecha_neg"] = datetime.strptime(p["fecha_neg"], "%Y-%m-%d").date()
                    p["fecha_ent"] = datetime.strptime(p["fecha_ent"], "%Y-%m-%d").date()
            except: pedidos = []
    else: pedidos = []

    if os.path.exists(ARCHIVO_ENTREGADOS):
        with open(ARCHIVO_ENTREGADOS, "r") as f:
            try:
                entregados = json.load(f)
                for e in entregados:
                    e["fecha_neg"] = datetime.strptime(e["fecha_neg"], "%Y-%m-%d").date()
                    e["fecha_ent"] = datetime.strptime(e["fecha_ent"], "%Y-%m-%d").date()
            except: entregados = []
    else: entregados = []

    if os.path.exists(ARCHIVO_PROVEEDORES):
        with open(ARCHIVO_PROVEEDORES, "r") as f:
            try: 
                proveedores = json.load(f)
                for p in proveedores:
                    if "historial_listas" not in p: p["historial_listas"] = []
            except: proveedores = []
    else: proveedores = []
    
    return pedidos, entregados, proveedores

def guardar_pedidos():
    pedidos_serializables = []
    for p in st.session_state.pedidos:
        pedidos_serializables.append({
            "proveedor": p["proveedor"], "orden": p["orden"], "detalle": p.get("detalle", ""),
            "fecha_neg": p["fecha_neg"].strftime("%Y-%m-%d"), "fecha_ent": p["fecha_ent"].strftime("%Y-%m-%d"),
            "archivo_nota": p.get("archivo_nota", ""), "archivo_factura": p.get("archivo_factura", "")
        })
    with open(ARCHIVO_PEDIDOS, "w") as f: json.dump(pedidos_serializables, f, indent=4)

def guardar_entregados():
    ent_serializables = []
    for e in st.session_state.entregados:
        ent_serializables.append({
            "proveedor": e["proveedor"], "orden": e["orden"], "detalle": e.get("detalle", ""),
            "fecha_neg": e["fecha_neg"].strftime("%Y-%m-%d"), "fecha_ent": e["fecha_ent"].strftime("%Y-%m-%d"),
            "archivo_nota": e.get("archivo_nota", ""), "archivo_factura": e.get("archivo_factura", "")
        })
    with open(ARCHIVO_ENTREGADOS, "w") as f: json.dump(ent_serializables, f, indent=4)

def guardar_proveedores():
    with open(ARCHIVO_PROVEEDORES, "w") as f: json.dump(st.session_state.proveedores, f, indent=4)

# --- INICIALIZAR MEMORIA ---
if 'datos_cargados' not in st.session_state:
    pedidos_db, entregados_db, proveedores_db = cargar_datos()
    st.session_state.pedidos = pedidos_db
    st.session_state.entregados = entregados_db
    st.session_state.proveedores = proveedores_db
    st.session_state.datos_cargados = True

if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

# --- CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="DuoLogic IA | Administración", page_icon="📊", layout="wide")
st.title("📊 DuoLogic IA - Suite de Compras y Proveedores")

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("🔐 Acceso al Sistema")
    st.markdown("Ingresá la clave corporativa para habilitar los módulos.")
    
    clave_usuario = st.text_input("Contraseña:", type="password")
    if st.button("Ingresar", type="primary"):
        if clave_usuario == "eleven2026":
            st.session_state.autenticado = True
        else:
            st.session_state.autenticado = False
            st.error("❌ Contraseña incorrecta")
    
    API_KEY_OCULTA = st.secrets["GEMINI_API_KEY"]
    
    if st.session_state.autenticado:
        api_key_input = API_KEY_OCULTA
        st.success("✅ Sistema Desbloqueado")
    else:
        api_key_input = None
            
    st.markdown("---")
    st.markdown("**Módulos Activos:**\n- 🧾 1. Extractor\n- 🤝 2. Negociación\n- ⚖️ 3. Comparador\n- 🚨 4. Reclamos\n- 🚚 5. Seguimiento\n- 📂 6. Proveedores\n- 📈 7. Análisis de Costos")

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(["🧾 Extractor", "🤝 Negociación", "⚖️ Comparador", "🚨 Reclamos", "🚚 Seguimiento", "📂 Proveedores", "📈 Análisis de Costos"])

# ==========================================
# MÓDULOS DE IA (1 a 4) 
# ==========================================
with tab1:
    st.header("Extractor Inteligente")
    if not st.session_state.autenticado: st.warning("⚠️ Módulo bloqueado.")
    else:
        texto_presupuesto = st.text_area("Texto del presupuesto:", height=100)
        if st.button("Extraer Datos a Tabla", type="primary"):
            with st.spinner('Analizando...'):
                try:
                    client = genai.Client(api_key=api_key_input)
                    st.markdown(client.models.generate_content(model='gemini-2.5-flash', contents=f"Actúa como Analista de Compras. Extrae en TABLA Markdown: Ítem, Cantidad, Precio Unitario, Subtotal. TEXTO: {texto_presupuesto}").text.replace('$', '\$'))
                except Exception as e: st.error(f"Error: {e}")

with tab2:
    st.header("Generador de Negociación")
    if not st.session_state.autenticado: st.warning("⚠️ Módulo bloqueado.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            proveedor = st.text_input("Proveedor:")
            servicio = st.text_input("Servicio/Producto:")
        with col2:
            precio_recibido = st.text_input("Precio cotizado:")
            precio_objetivo = st.text_input("Precio objetivo:")
        tono = st.selectbox("Tono:", ["Cordial y colaborativo", "Firme", "Descuento"])
        if st.button("Generar Correo", type="primary"):
            with st.spinner('Redactando...'):
                try:
                    client = genai.Client(api_key=api_key_input)
                    st.markdown(client.models.generate_content(model='gemini-2.5-flash', contents=f"Redacta correo para negociar. Proveedor: {proveedor}. Servicio: {servicio}. Cotizado: {precio_recibido}. Objetivo: {precio_objetivo}. Tono: {tono}.").text.replace('$', '\$'))
                except Exception as e: st.error(f"Error: {e}")

with tab3:
    st.header("⚖️ Comparador de Presupuestos")
    if not st.session_state.autenticado: st.warning("⚠️ Módulo bloqueado.")
    else:
        col_a, col_b, col_c = st.columns(3)
        with col_a: cotizacion_a = st.text_area("Cotización A:", height=100)
        with col_b: cotizacion_b = st.text_area("Cotización B:", height=100)
        with col_c: cotizacion_c = st.text_area("Cotización C:", height=100)
        if st.button("Comparar Cotizaciones", type="primary"):
            with st.spinner('Cruzando datos...'):
                try:
                    client = genai.Client(api_key=api_key_input)
                    st.markdown(client.models.generate_content(model='gemini-2.5-flash', contents=f"Analiza: Cot A: {cotizacion_a}, Cot B: {cotizacion_b}, Cot C: {cotizacion_c}. NO uses el símbolo $, usa USD o ARS.").text.replace('$', '\$'))
                except Exception as e: st.error(f"Error: {e}")

with tab4:
    st.header("🚨 Gestión de Reclamos")
    if not st.session_state.autenticado: st.warning("⚠️ Módulo bloqueado.")
    else:
        col3, col4 = st.columns(2)
        with col3:
            prov_reclamo = st.text_input("Proveedor a reclamar:")
            motivo = st.selectbox("Motivo:", ["Factura mal", "Demora", "Defecto"])
        with col4:
            detalle = st.text_area("Detalle:")
            solucion = st.text_input("Solución:")
        if st.button("Redactar Reclamo", type="primary"):
            with st.spinner('Redactando...'):
                try:
                    client = genai.Client(api_key=api_key_input)
                    st.markdown(client.models.generate_content(model='gemini-2.5-flash', contents=f"Escribe correo de reclamo a {prov_reclamo}. Motivo: {motivo}. Detalle: {detalle}. Solución: {solucion}.").text.replace('$', '\$'))
                except Exception as e: st.error(f"Error: {e}")

# ==========================================
# PESTAÑA 5: SEGUIMIENTO
# ==========================================
with tab5:
    st.header("🚚 Panel de Seguimiento y Control de Pedidos")
    if not st.session_state.autenticado: st.warning("⚠️ Módulo bloqueado.")
    else:
        hoy = date.today()
        with st.expander("➕ Cargar Nuevo Pedido con Archivos", expanded=False):
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                lista_nombres_prov = [p["nombre"] for p in st.session_state.proveedores]
                opciones_prov = ["-- Seleccionar Proveedor --"] + lista_nombres_prov + ["➕ Ingresar nuevo manualmente..."]
                seleccion_prov = st.selectbox("Proveedor (Buscá o seleccioná uno existente):", opciones_prov)
                
                if seleccion_prov == "➕ Ingresar nuevo manualmente...": nuevo_prov = st.text_input("Escribí el nombre del nuevo proveedor:")
                elif seleccion_prov == "-- Seleccionar Proveedor --": nuevo_prov = ""
                else: nuevo_prov = seleccion_prov
                
                nueva_orden = st.text_input("Nº de Orden / Referencia:")
                detalle_ped = st.text_area("Detalle del pedido:")
            
            with col_p2:
                fecha_neg = st.date_input("Fecha de Negociación:", value=hoy)
                fecha_ent = st.date_input("Fecha de Entrega Pactada:", value=hoy)
                archivo_nota = st.file_uploader("Adjuntar Nota de Pedido", type=["pdf", "png", "jpg"], key="up_nota")
                archivo_factura = st.file_uploader("Adjuntar Factura", type=["pdf", "png", "jpg"], key="up_factura")
            
            if st.button("Guardar Pedido", type="primary"):
                if nuevo_prov and nueva_orden:
                    ruta_nota, ruta_fact = "", ""
                    if archivo_nota:
                        ruta_nota = os.path.join(CARPETA_PEDIDOS, f"Nota_{nueva_orden}_{archivo_nota.name}")
                        with open(ruta_nota, "wb") as f: f.write(archivo_nota.getbuffer())
                    if archivo_factura:
                        ruta_fact = os.path.join(CARPETA_PEDIDOS, f"Factura_{nueva_orden}_{archivo_factura.name}")
                        with open(ruta_fact, "wb") as f: f.write(archivo_factura.getbuffer())
                        
                    proveedor_existe = False
                    for prov in st.session_state.proveedores:
                        if prov["nombre"].strip().lower() == nuevo_prov.strip().lower():
                            proveedor_existe = True
                            if "historial_facturas" not in prov: prov["historial_facturas"] = []
                            if ruta_fact:
                                prov["historial_facturas"].insert(0, {"orden": nueva_orden, "fecha": fecha_ent.strftime('%d/%m/%Y'), "ruta": ruta_fact})
                                prov["historial_facturas"] = prov["historial_facturas"][:3]
                            break
                    
                    if not proveedor_existe:
                        nuevo_historial = [{"orden": nueva_orden, "fecha": fecha_ent.strftime('%d/%m/%Y'), "ruta": ruta_fact}] if ruta_fact else []
                        st.session_state.proveedores.append({
                            "nombre": nuevo_prov, "rubro": "Autogenerado", "contacto": "-", "lista_precios": "",
                            "historial_facturas": nuevo_historial, "historial_listas": []
                        })
                    guardar_proveedores()

                    st.session_state.pedidos.append({
                        "proveedor": nuevo_prov, "orden": nueva_orden, "detalle": detalle_ped,
                        "fecha_neg": fecha_neg, "fecha_ent": fecha_ent,
                        "archivo_nota": ruta_nota, "archivo_factura": ruta_fact
                    })
                    guardar_pedidos()
                    st.success("¡Pedido y factura guardados!")
                    st.rerun()
                else: st.warning("Completa Proveedor y Orden.")

        st.markdown("---")
        st.subheader("📋 Estado de Pedidos Activos")
        if not st.session_state.pedidos: st.info("Todo al día ✅")
        else:
            for index, ped in enumerate(st.session_state.pedidos):
                diferencia_dias = (hoy - ped['fecha_ent']).days
                with st.container():
                    c1, c2, c3, c4 = st.columns([2, 2, 2, 3])
                    c1.write(f"🏢 **{ped['proveedor']}** (Ord: {ped['orden']})")
                    c2.write(f"📦 Entrega: {ped['fecha_ent'].strftime('%d/%m/%Y')}")
                    if diferencia_dias > 0: c3.error(f"⚠️ Atrasado: {diferencia_dias} días")
                    elif diferencia_dias == 0: c3.warning("⏳ HOY")
                    else: c3.success(f"✅ Faltan {abs(diferencia_dias)} días")
                    with c4:
                        col_b1, col_b2 = st.columns(2)
                        if col_b1.button("✅ Entregado", key=f"ok_{index}"):
                            st.session_state.entregados.append(st.session_state.pedidos.pop(index))
                            guardar_pedidos()
                            guardar_entregados()
                            st.rerun()
                        if col_b2.button("🗑️", key=f"del_{index}"):
                            st.session_state.pedidos.pop(index)
                            guardar_pedidos()
                            st.rerun()
                st.markdown("---")

        st.markdown("### 📦 Historial Entregados")
        busqueda = st.text_input("🔎 Buscar historial:")
        if st.session_state.entregados:
            resultados = [e for e in st.session_state.entregados if not busqueda or busqueda.lower() in e['proveedor'].lower() or busqueda.lower() in e['orden'].lower()]
            for index, ent in enumerate(resultados):
                with st.expander(f"✅ {ent['proveedor']} - Orden: {ent['orden']}"):
                    d1, d2 = st.columns(2)
                    if ent.get('archivo_nota') and os.path.exists(ent['archivo_nota']):
                        with open(ent['archivo_nota'], "rb") as file: d1.download_button("📄 Nota", file, os.path.basename(ent['archivo_nota']), key=f"dn_{index}")
                    if ent.get('archivo_factura') and os.path.exists(ent['archivo_factura']):
                        with open(ent['archivo_factura'], "rb") as file: d2.download_button("🧾 Factura", file, os.path.basename(ent['archivo_factura']), key=f"df_{index}")

# ==========================================
# PESTAÑA 6: BASE DE PROVEEDORES
# ==========================================
with tab6:
    st.header("📂 Directorio y Base de Proveedores")
    if not st.session_state.autenticado: st.warning("⚠️ Módulo bloqueado.")
    else:
        with st.expander("➕ Cargar Nuevo Proveedor", expanded=False):
            col_prov1, col_prov2 = st.columns(2)
            with col_prov1:
                prov_nombre = st.text_input("Nombre:")
                prov_rubro = st.text_input("Rubro / Servicio:")
                prov_contacto = st.text_input("Email o Teléfono:")
            with col_prov2:
                prov_lista = st.text_input("Link web a Lista de Precios:")
                archivo_lista_inicial = st.file_uploader("Adjuntar Lista Inicial (Excel o PDF)", type=["pdf", "xlsx", "xls", "csv"], key="up_lis_ini")
                fecha_lista_inicial = st.date_input("Fecha de emisión de esta lista:", value=date.today(), key="fecha_ini")
                
            if st.button("Guardar Proveedor", type="primary"):
                if prov_nombre and prov_rubro:
                    nuevo_historial_listas = []
                    if archivo_lista_inicial:
                        nombre_seguro = f"Lista_{prov_nombre}_{fecha_lista_inicial.strftime('%Y%m%d')}_{archivo_lista_inicial.name}"
                        ruta_lista = os.path.join(CARPETA_PROVEEDORES, nombre_seguro)
                        with open(ruta_lista, "wb") as f: f.write(archivo_lista_inicial.getbuffer())
                        nuevo_historial_listas.append({"fecha": fecha_lista_inicial.strftime('%d/%m/%Y'), "ruta": ruta_lista})

                    st.session_state.proveedores.append({
                        "nombre": prov_nombre, "rubro": prov_rubro, "contacto": prov_contacto, 
                        "lista_precios": prov_lista, "historial_facturas": [], "historial_listas": nuevo_historial_listas
                    })
                    guardar_proveedores()
                    st.rerun()
                else:
                    st.warning("El Nombre y Rubro son obligatorios.")

        st.markdown("---")
        busqueda_prov = st.text_input("🔎 Buscar proveedor:")
        proveedores_filtrados = st.session_state.proveedores
        if busqueda_prov:
            proveedores_filtrados = [p for p in st.session_state.proveedores if busqueda_prov.lower() in p['nombre'].lower() or busqueda_prov.lower() in p['rubro'].lower()]
        
        if not proveedores_filtrados: st.info("No se encontraron proveedores.")
        else:
            for prov in proveedores_filtrados:
                original_index = st.session_state.proveedores.index(prov)
                with st.expander(f"🏢 {prov['nombre']} | 🏷️ {prov['rubro']}", expanded=False):
                    c_info1, c_info2 = st.columns(2)
                    c_info1.markdown(f"**📞 Contacto:** {prov['contacto']}")
                    if prov.get('lista_precios'): c_info2.markdown(f"**🔗 Link Web:** [Ver Online]({prov['lista_precios']})")
                    
                    st.markdown("---")
                    col_hist1, col_hist2 = st.columns(2)
                    
                    with col_hist1:
                        st.markdown("📦 **Últimas Facturas:**")
                        historial_fac = prov.get("historial_facturas", [])
                        if not historial_fac: st.caption("Sin facturas.")
                        else:
                            for i, fac in enumerate(historial_fac):
                                if os.path.exists(fac["ruta"]):
                                    with open(fac["ruta"], "rb") as f: st.download_button(label=f"🧾 Ord {fac['orden']} ({fac['fecha']})", data=f, file_name=os.path.basename(fac["ruta"]), key=f"dl_fac_{original_index}_{i}")
                                else: st.error("⚠️ Borrado")

                    with col_hist2:
                        st.markdown("📋 **Últimas Listas de Precios:**")
                        historial_lis = prov.get("historial_listas", [])
                        if not historial_lis: st.caption("Sin listas.")
                        else:
                            for i, lis in enumerate(historial_lis):
                                c_btn1, c_btn2 = st.columns([4, 1])
                                with c_btn1:
                                    if os.path.exists(lis["ruta"]):
                                        with open(lis["ruta"], "rb") as f: st.download_button(label=f"⬇️ {lis['fecha']}", data=f, file_name=os.path.basename(lis["ruta"]), key=f"dl_lis_{original_index}_{i}")
                                    else: st.error("⚠️ Borrado")
                                with c_btn2:
                                    if st.button("❌", key=f"del_lis_btn_{original_index}_{i}", help="Eliminar este archivo"):
                                        st.session_state.proveedores[original_index]["historial_listas"].pop(i)
                                        guardar_proveedores()
                                        st.rerun()
                        
                        st.markdown("**➕ Subir nueva Lista:**")
                        nueva_lista = st.file_uploader("Adjuntar Excel o PDF", type=["pdf", "xlsx", "xls", "csv"], key=f"up_lista_{original_index}")
                        fecha_nueva_lista = st.date_input("Fecha de emisión del archivo:", value=date.today(), key=f"d_lista_{original_index}")
                        
                        if st.button("Guardar Lista", key=f"btn_lista_{original_index}"):
                            if nueva_lista:
                                if "historial_listas" not in prov: prov["historial_listas"] = []
                                nombre_seguro = f"Lista_{prov['nombre']}_{fecha_nueva_lista.strftime('%Y%m%d')}_{nueva_lista.name}"
                                ruta_lista = os.path.join(CARPETA_PROVEEDORES, nombre_seguro)
                                with open(ruta_lista, "wb") as f: f.write(nueva_lista.getbuffer())
                                prov["historial_listas"].insert(0, {"fecha": fecha_nueva_lista.strftime('%d/%m/%Y'), "ruta": ruta_lista})
                                prov["historial_listas"] = prov["historial_listas"][:3]
                                guardar_proveedores()
                                st.success("Lista guardada correctamente.")
                                st.rerun()
                            else: st.warning("Seleccioná un archivo.")
                    st.write("") 
                    if st.button("🗑️ Eliminar Proveedor completo", key=f"delprov_{original_index}"):
                        st.session_state.proveedores.pop(original_index)
                        guardar_proveedores()
                        st.rerun()

# ==========================================
# PESTAÑA 7: ANÁLISIS HISTÓRICO Y CREADOR DE WORD
# ==========================================
with tab7:
    st.header("📈 Análisis Histórico de Precios Automático")
    if not st.session_state.autenticado:
        st.warning("⚠️ Módulo bloqueado. Ingresá la contraseña en el menú lateral.")
    else:
        st.markdown("Este módulo cruza las últimas dos listas de precios de un proveedor para calcular la inflación y variaciones.")
        
        proveedores_aptos = [p for p in st.session_state.proveedores if len(p.get("historial_listas", [])) >= 2]
        lista_nombres_aptos = [p["nombre"] for p in proveedores_aptos]
        
        if not lista_nombres_aptos:
            st.info("📌 Subí al menos **2 Listas de Precios** diferentes a un mismo proveedor para usar esta función.")
        else:
            prov_elegido = st.selectbox("Seleccionar Proveedor para auditar:", ["-- Seleccionar --"] + lista_nombres_aptos)
            
            if prov_elegido != "-- Seleccionar --":
                prov_data = next(p for p in proveedores_aptos if p["nombre"] == prov_elegido)
                listas = prov_data["historial_listas"]
                
                st.write(f"🔄 **Comparando:** Lista antigua ({listas[1]['fecha']}) 🆚 Lista nueva ({listas[0]['fecha']})")
                
                if st.button("📊 Generar Informe de Aumentos", type="primary"):
                    with st.spinner("Leyendo documentos y cruzando datos. Esto puede demorar un minuto..."):
                        try:
                            client = genai.Client(api_key=api_key_input)
                            
                            def procesar_archivo_para_gemini(ruta):
                                if ruta.endswith(('.xlsx', '.xls')):
                                    df = pd.read_excel(ruta)
                                    return df.to_csv(index=False)
                                elif ruta.endswith('.csv'):
                                    df = pd.read_csv(ruta)
                                    return df.to_csv(index=False)
                                else:
                                    return client.files.upload(file=ruta)

                            archivo_viejo = procesar_archivo_para_gemini(listas[1]["ruta"])
                            archivo_nuevo = procesar_archivo_para_gemini(listas[0]["ruta"])
                            
                            prompt = f"""Actúa como un Gerente de Finanzas y Control de Costos. 
                            Tienes dos listas de precios del proveedor {prov_elegido}.
                            El Archivo 1 es del {listas[1]['fecha']} (Antiguo).
                            El Archivo 2 es del {listas[0]['fecha']} (Nuevo).
                            
                            Realiza un informe ejecutivo detallando:
                            1. El porcentaje de aumento promedio estimado general.
                            2. Tabla Markdown con los 5 productos que más aumentaron (Nombre, Precio Viejo, Precio Nuevo, % Aumento). Usa formato de tabla Markdown estandar.
                            3. Conclusión sobre impacto y sugerencia de renegociación.
                            """
                            
                            respuesta = client.models.generate_content(
                                model='gemini-2.5-flash',
                                contents=[prompt, "ARCHIVO ANTIGUO:", archivo_viejo, "ARCHIVO NUEVO:", archivo_nuevo]
                            )
                            st.success("✅ Análisis Completado")
                            st.markdown(respuesta.text.replace('$', '\$'))
                            
                            # --- MAGIA: CREACIÓN DEL ARCHIVO WORD CON TABLAS REALES ---
                            doc = Document()
                            doc.add_heading(f'Informe Ejecutivo de Costos - {prov_elegido}', level=1)
                            
                            texto_limpio = respuesta.text.replace('**', '') 
                            lineas = texto_limpio.split('\n')
                            
                            en_tabla = False
                            tabla_word = None
                            
                            for linea in lineas:
                                linea_limpia = linea.strip()
                                if not linea_limpia:
                                    continue
                                
                                # Si detectamos que la línea es parte de una tabla Markdown
                                if '|' in linea_limpia:
                                    # Ignoramos la línea divisoria del Markdown (ej: |---|---|)
                                    if '---' in linea_limpia:
                                        continue
                                        
                                    # Separamos las celdas
                                    celdas = [c.strip() for c in linea_limpia.split('|')]
                                    # Limpiamos elementos vacíos si la tabla tiene bordes exteriores |
                                    if celdas and celdas[0] == '': celdas.pop(0)
                                    if celdas and celdas[-1] == '': celdas.pop()
                                    
                                    if not en_tabla:
                                        # Creamos la tabla en Word
                                        en_tabla = True
                                        tabla_word = doc.add_table(rows=1, cols=len(celdas))
                                        tabla_word.style = 'Table Grid'
                                        hdr_cells = tabla_word.rows[0].cells
                                        for i, texto_celda in enumerate(celdas):
                                            if i < len(hdr_cells):
                                                hdr_cells[i].text = texto_celda
                                    else:
                                        # Añadimos fila a la tabla de Word
                                        row_cells = tabla_word.add_row().cells
                                        for i, texto_celda in enumerate(celdas):
                                            if i < len(row_cells):
                                                row_cells[i].text = texto_celda
                                else:
                                    en_tabla = False
                                    doc.add_paragraph(linea_limpia)
                                    
                            buffer = io.BytesIO()
                            doc.save(buffer)
                            buffer.seek(0)
                            
                            st.download_button(
                                label="📄 Descargar Informe en Word (.docx)",
                                data=buffer,
                                file_name=f"Informe_Aumentos_{prov_elegido.replace(' ', '_')}_{date.today().strftime('%Y%m%d')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                        except Exception as e:
                            st.error(f"❌ Error al generar el documento. Detalle técnico: {e}")